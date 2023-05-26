# coding:utf-8

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches
import docx

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


class BaseValue(object):

    def __init__(self, name=''):
        self.name = name

    def __str__(self):
        return self.name

    __repr__ = __str__


class TableFile(BaseValue):

    def __init__(self, name=''):
        super(TableFile, self).__init__(name)
        self.tables = []

    def add_table(self, table):
        self.tables.append(table)

    def __str__(self):
        return '\n'.join(t for t in self.tables)
        # return ''.join(t for t in self.tables)

    __repr__ = __str__


class Table(BaseValue):

    def __init__(self, name=''):
        super(Table, self).__init__(name)
        self.modules = []

    def add_module(self, module):
        self.modules.append(module)

    def __str__(self):
        s = ''
        s = s + self.name + '\n'
        s = s + '\n'.join(str(t) for t in self.modules)
        return s

    __repr__ = __str__


class BigModule(BaseValue):

    def __init__(self, name=''):
        super(BigModule, self).__init__(name)
        self.level = 1
        self.small_module = []

    def add_small_module(self, small_module):
        self.small_module.append(small_module)

    def __str__(self):
        s = ''
        s = s + self.name + '\n'
        s = s + '\n'.join(str(t) for t in self.small_module)
        return s

    __repr__ = __str__


class SmallModule(BaseValue):

    def __init__(self, name=''):
        super(SmallModule, self).__init__(name)
        self.name = name
        self.level = 2
        self.cases = []

    def add_case(self, case):
        self.cases.append(case)

    def __str__(self):
        s = ''
        s = s + self.name + '\n'
        s = s + '\n'.join(str(t) for t in self.cases)
        return s

    __repr__ = __str__


class Case(BaseValue):

    def __init__(self, name=''):
        super(Case, self).__init__(name)
        self.level = 3
        self.premise = []
        self.step = []
        self.hope = []
        self.now_step = 0
        self.now_premise = 0

    def add_premise(self, premise):
        if pd.isna(premise):
            premise = ''
        if premise.strip():
            self.now_premise = self.now_premise + 1
            premise = "{}、{}".format(self.now_premise, premise)
            self.premise.append(premise)

    def add_step(self, step):
        if pd.isna(step):
            step = ''
        self.now_step += 1
        step = "{}、{}".format(self.now_step, step)
        self.step.append(step)

    def add_hope(self, hope):
        if pd.isna(hope):
            hope = ''
        hope = "{}、{}".format(self.now_step, hope)
        self.hope.append(hope)

    def __str__(self):
        s = ''
        s = s + self.name + '\n'
        s = s + 'premise\n' + ' | '.join(self.premise)
        s = s + '\n'
        s = s + 'step\n' + ' | '.join(self.step)
        s = s + '\n'
        s = s + 'hope\n' + ' | '.join(self.hope)
        s = s + '\n'
        return s

    __repr__ = __str__


def deal_with_title(title):
    names = title.split('、')
    title_id = names[0]
    name = names[1]
    count = title_id.count('_')
    return count, name


def is_big_module(title):
    if not pd.isna(title):
        names = title.split('、')
        title_id = names[0]
        count = title_id.count('_')
        if count == 3:
            return True
    return False


def is_small_module(title):
    if not pd.isna(title):
        names = title.split('、')
        title_id = names[0]
        count = title_id.count('_')
        if count == 4:
            return True
    return False


if __name__ == '__main__':
    case_file_path = u"/home/frz/桌面/C21-HV3.0-测试用例修改版.xlsx"
    sheet_name = u"F1_1_1、有线网络"
    save_path = u"/home/frz/桌面/测试.docx"
    head_name = "生产测试"
    tf = TableFile('用例文件')
    tb = Table('用例表格')
    df = pd.read_excel(case_file_path, sheet_name)
    now_bm = BigModule()
    now_sm = SmallModule()
    now_case = Case()
    tmp_case_name = ''
    flag = 0
    for value in df.values[5::]:
        title = value[0]
        point = value[1]
        premise = value[2]
        step = value[4]
        hope = value[5]
        if not pd.isna(title):
            if flag == 1:
                now_case.name = tmp_case_name
                now_sm.add_case(now_case)
                flag = 0
            count, name = deal_with_title(title)
            if count == 3:
                now_bm = BigModule()
                now_bm.name = name
                tb.add_module(now_bm)
            elif count == 4:
                now_sm = SmallModule()
                now_sm.name = name
                now_bm.add_small_module(now_sm)
            elif count == 5:
                now_case = Case()
                tmp_case_name = name
                flag = 1
        else:
            if not pd.isna(point):
                flag = 0
                now_case = Case()
                now_case.name = tmp_case_name + '-' + point
                now_sm.add_case(now_case)
            now_case.add_premise(premise)
            now_case.add_step(step)
            now_case.add_hope(hope)
    # print(tb)

    # 打开文档
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 加入不同等级的标题
    # document.add_heading(head_name, 0)
    head = document.add_heading("", level=0)  # 这里不填标题内容
    run = head.add_run(head_name)
    run.font.name = u'Cambria'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0, 0, 0)

    # 插入修订记录
    document.add_paragraph(u'')
    document.add_paragraph(u'修订记录')
    table = document.add_table(rows=3, cols=4, style='Table Grid')
    table.cell(0, 0).text = u"修订版本"
    table.cell(0, 1).text = u"修订日期"
    table.cell(0, 2).text = u"修订人"
    table.cell(0, 3).text = u"修订内容"
    table.cell(0, 3).width = Inches(4)
    document.add_paragraph(u'')

    bm_id = 0
    for bm in tb.modules:
        bm_id += 1
        sm_id = 2
        # document.add_heading(u'{}.{}'.format(bm_id, bm.name), 1)
        head = document.add_heading("", level=1)  # 这里不填标题内容
        run = head.add_run(u'{} {}'.format(bm_id, bm.name))
        run.font.name = u'Cambria'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
        run.font.color.rgb = RGBColor(0, 0, 0)

        head = document.add_heading("", level=2)  # 这里不填标题内容
        run = head.add_run(u'{}.{} {}'.format(bm_id, 1, u'需求分析'))
        run.font.name = u'Cambria'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
        run.font.color.rgb = RGBColor(0, 0, 0)
        document.add_paragraph(u'')

        head = document.add_heading("", level=2)  # 这里不填标题内容
        run = head.add_run(u'{}.{} {}'.format(bm_id, 2, u'功能说明'))
        run.font.name = u'Cambria'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
        run.font.color.rgb = RGBColor(0, 0, 0)
        document.add_paragraph(u'')

        for sm in bm.small_module:
            sm_id += 1
            case_id = 0

            # document.add_heading(u'{}.{}.{}'.format(bm_id, sm_id, sm.name), 2)
            head = document.add_heading("", level=2)  # 这里不填标题内容
            run = head.add_run(u'{}.{} {}'.format(bm_id, sm_id, sm.name))
            run.font.name = u'Cambria'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
            run.font.color.rgb = RGBColor(0, 0, 0)
            for case in sm.cases:
                case_id += 1
                # document.add_heading(u'{}.{}.{}.{}'.format(bm_id, sm_id, case_id, case.name), 3)
                head = document.add_heading("", level=3)  # 这里不填标题内容
                run = head.add_run(u'{}.{}.{} {}'.format(bm_id, sm_id, case_id, case.name))
                run.font.name = u'Cambria'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
                run.font.color.rgb = RGBColor(0, 0, 0)
                document.add_paragraph(u'')
                table = document.add_table(rows=5, cols=4, style='Table Grid')
                table.cell(0, 1).merge(table.cell(0, 3))
                table.cell(1, 1).merge(table.cell(1, 3))
                table.cell(2, 1).merge(table.cell(2, 3))
                table.cell(3, 0).merge(table.cell(3, 1))
                table.cell(3, 2).merge(table.cell(3, 3))
                table.cell(4, 0).merge(table.cell(4, 1))
                table.cell(4, 2).merge(table.cell(4, 3))
                table.cell(0, 0).text = u"测试项"
                table.cell(0, 1).text = u"{}".format(case.name)
                table.cell(1, 0).text = u"前置条件"
                table.cell(1, 1).text = u'\n'.join(case.premise).strip()
                table.cell(2, 0).text = u"其他说明"
                table.cell(2, 1).text = u""
                table.cell(3, 0).text = u"测试步骤"
                table.cell(4, 0).text = u'\n'.join(case.step).strip()
                table.cell(3, 2).text = u"预期结果"
                table.cell(4, 2).text = u'\n'.join(case.hope).strip()
                document.add_paragraph(u'')
    # 保存文件
    document.save(save_path)
