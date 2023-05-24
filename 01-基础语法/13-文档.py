from docx import Document
from docx.shared import Inches

# 创建一个新的Word文档
document = Document()

# 添加标题
document.add_heading('Document Title', 0)

# 添加段落
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# 添加图片
document.add_picture('image.png', width=Inches(1.25))

# 添加表格
table = document.add_table(rows=2, cols=2)
cell = table.cell(0, 1)
cell.text = 'This is a cell'

# 保存文档
document.save('demo.docx')